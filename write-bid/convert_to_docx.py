"""
将 Markdown 技术方案转换为 Word 文档（完整版）。
修复原版问题：
1. 不丢失内容（原版跳过空行导致段落丢失）
2. 正确处理标题、列表、表格
3. 设置正确的段落间距和缩进
4. 输出字数统计，避免字数不足
"""
import sys, re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
except ImportError as e:
    print(f"请先安装: pip install python-docx （{e}）", file=sys.stderr)
    sys.exit(1)

if len(sys.argv) < 2:
    print("用法: python convert_to_docx.py <方案.md> [输出文件.docx]", file=sys.stderr)
    sys.exit(1)

md_path = Path(sys.argv[1])
if not md_path.exists():
    print(f"文件不存在: {md_path}", file=sys.stderr)
    sys.exit(1)

out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix('.docx')

doc = Document()

# 设置页边距
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

# 设置默认样式
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_run_font(run, name='宋体', size=12, bold=False):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_heading_text(doc, text, level):
    h = doc.add_heading(level=level)
    h.clear()
    run = h.add_run(text)
    sizes = {0: 22, 1: 18, 2: 15, 3: 13, 4: 12}
    set_run_font(run, '黑体', sizes.get(level, 12), bold=True)
    h.paragraph_format.space_before = Pt(12)
    h.paragraph_format.space_after = Pt(6)
    return h

def add_normal_paragraph(doc, text):
    if not text.strip():
        return
    para = doc.add_paragraph()
    para.paragraph_format.first_line_indent = Pt(24)
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(3)
    run = para.add_run(text.strip())
    set_run_font(run, '宋体', 12)
    return para

def add_table_from_lines(doc, table_lines):
    rows = []
    for line in table_lines:
        s = line.strip()
        if not s or not s.startswith('|'):
            continue
        if re.match(r'^\|[-:| ]+\|$', s):
            continue
        cells = [c.strip() for c in s.strip('|').split('|')]
        if cells:
            rows.append(cells)
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=ncols)
    table.style = 'Table Grid'
    table.autofit = True
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= ncols:
                break
            cell = table.cell(i, j)
            cell.text = ''
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.font.size = Pt(9)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if i == 0:
                run.bold = True
    doc.add_paragraph()

md = md_path.read_text(encoding='utf-8')
lines = md.split('\n')

total_chars = 0
para_count = 0
table_count = 0

i = 0
while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    if not stripped:
        i += 1
        continue

    # 标题
    m = re.match(r'^(#{1,4})\s+(.*)', line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        add_heading_text(doc, text, level)
        i += 1
        continue

    # 表格
    if stripped.startswith('|') and stripped.endswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i])
            i += 1
        add_table_from_lines(doc, table_lines)
        table_count += 1
        continue

    # 列表项
    if stripped.startswith('- ') or stripped.startswith('* '):
        text = re.sub(r'^[-*]\s+', '', stripped)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run('• ' + text)
        set_run_font(run, '宋体', 12)
        total_chars += len(text)
        i += 1
        continue

    # 数字列表
    if re.match(r'^\d+[.、]\s', stripped):
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', stripped)
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1)
        para.paragraph_format.line_spacing = 1.5
        para.paragraph_format.space_after = Pt(3)
        run = para.add_run(text)
        set_run_font(run, '宋体', 12)
        total_chars += len(text)
        i += 1
        continue

    # 分隔线
    if stripped in ('---', '___', '***'):
        i += 1
        continue

    # 正文段落
    text = stripped
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    add_normal_paragraph(doc, text)
    total_chars += len(text)
    para_count += 1
    i += 1

doc.save(str(out_path))

print(f"已导出: {out_path}")
print(f"原Markdown字符数: {len(md)}")
print(f"写入段落数: {para_count}")
print(f"写入表格数: {table_count}")
print(f"正文实际写入字符数: {total_chars}")
print(f"估算Word字数: ~{total_chars} 字")
if total_chars < 96000:
    print(f"[警告] 字数不足目标（96,000字），还差 {96000 - total_chars} 字，请补充内容后重新导出")
else:
    print(f"[达标] 字数达到目标要求（>=96,000字）")